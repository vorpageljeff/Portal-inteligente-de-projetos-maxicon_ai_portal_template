import io
from collections.abc import Generator

import pytest
from app.core.config import settings
from app.db.base import Base
from app.db.session import get_db
from app.main import app
from app.models import *  # noqa: F403
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker
from sqlalchemy.pool import StaticPool


@pytest.fixture()
def client(monkeypatch: pytest.MonkeyPatch) -> Generator[TestClient, None, None]:
    monkeypatch.setattr(settings, "ai_provider", "mock")
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    testing_session = sessionmaker(bind=engine, autoflush=False, autocommit=False)
    Base.metadata.create_all(bind=engine)

    def override_get_db() -> Generator[Session, None, None]:
        with testing_session() as db:
            yield db

    app.dependency_overrides[get_db] = override_get_db
    yield TestClient(app)
    app.dependency_overrides.clear()


def authenticate(client: TestClient) -> tuple[dict[str, str], str]:
    created = client.post(
        "/api/v1/auth/bootstrap-admin",
        json={
            "email": "lpn-admin@maxicon.com.br",
            "full_name": "Admin LPN",
            "password": "senha-forte-123",
        },
    )
    assert created.status_code == 201
    login = client.post(
        "/api/v1/auth/login",
        json={"email": "lpn-admin@maxicon.com.br", "password": "senha-forte-123"},
    )
    return {"Authorization": f"Bearer {login.json()['access_token']}"}, created.json()["id"]


def test_complete_lpn_flow_with_approval_documents_and_clone(client: TestClient) -> None:
    headers, user_id = authenticate(client)
    organization = client.get("/api/v1/organizations", headers=headers).json()[0]
    headers["X-Organization-ID"] = organization["id"]

    created_client = client.post(
        "/api/v1/organizations/clients",
        headers=headers,
        json={"name": "Cooperativa Exemplo"},
    )
    assert created_client.status_code == 201
    client_id = created_client.json()["id"]

    demand = client.post(
        "/api/v1/lpns/demands",
        headers=headers,
        json={
            "client_id": client_id,
            "title": "Automatizar recebimento",
            "external_number": "238862",
            "business_area": "Compras",
            "business_process": "Recebimento",
            "system_product": "ERP",
            "requester_name": "Ana Souza",
            "product_owner_name": "Jefferson Vorpagel",
            "priority": "high",
            "priority_reason": "Risco operacional relevante",
            "discovery_date": "2026-08-24",
            "demand_type": "improvement",
        },
    )
    assert demand.status_code == 201
    lpn_response = client.post(
        f"/api/v1/lpns/from-demand/{demand.json()['id']}", headers=headers
    )
    assert lpn_response.status_code == 201
    lpn = lpn_response.json()
    version_id = lpn["current_version"]["id"]

    item_ids: dict[str, str] = {}
    for kind, code, title, payload in (
        (
            "storytelling",
            "ST-001",
            "Contexto da demanda",
            {
                "actor": "Conferente",
                "current_situation": "Recebimento manual",
                "problem": "Duplicidade",
                "consequence": "Retrabalho",
                "desired_result": "Validar automaticamente",
            },
        ),
        ("stakeholder", "SK-001", "Conferente", {"responsibility": "Conferir material"}),
        ("requirement", "RF-001", "Validar quantidade", {"description": "Validar limite"}),
        (
            "acceptance_criterion",
            "CA-001",
            "Bloqueio acima do limite",
            {"given": "Pedido de 100", "when": "Informar 101", "then": "Bloquear"},
        ),
    ):
        response = client.post(
            f"/api/v1/lpns/versions/{version_id}/content",
            headers=headers,
            json={"kind": kind, "code": code, "title": title, "payload": payload},
        )
        assert response.status_code == 201
        item_ids[kind] = response.json()["id"]

    link = client.post(
        f"/api/v1/lpns/versions/{version_id}/links",
        headers=headers,
        json={
            "source_item_id": item_ids["requirement"],
            "target_item_id": item_ids["acceptance_criterion"],
            "relationship": "validated_by",
        },
    )
    assert link.status_code == 201

    for process_type in ("as_is", "to_be"):
        diagram = client.put(
            f"/api/v1/lpns/versions/{version_id}/diagrams",
            headers=headers,
            json={
                "process_type": process_type,
                "name": f"Processo {process_type}",
                "model": {
                    "lanes": [{"id": "lane-1", "name": "Conferente"}],
                    "nodes": [{"id": "start", "type": "start", "lane_id": "lane-1"}],
                    "edges": [],
                    "layout": {},
                },
            },
        )
        assert diagram.status_code == 200

    ai_preview = client.post(
        f"/api/v1/lpns/versions/{version_id}/ai/preview",
        headers=headers,
        json={
            "use_case": "objective",
            "input_text": (
                "O conferente precisa reduzir a duplicidade no recebimento e validar "
                "automaticamente a quantidade informada contra o pedido."
            ),
        },
    )
    assert ai_preview.status_code == 200
    suggestion_id = ai_preview.json()["suggestion_ids"][0]
    ai_decision = client.post(
        f"/api/v1/lpns/ai/suggestions/{suggestion_id}/decision",
        headers=headers,
        json={"decision": "accepted"},
    )
    assert ai_decision.status_code == 201

    approval = client.post(
        f"/api/v1/lpns/versions/{version_id}/approval",
        headers=headers,
        json={"approver_ids": [user_id], "required_approvals": 1},
    )
    assert approval.status_code == 201

    transitions = [
        "in_discovery",
        "as_is_validation",
        "as_is_approved",
        "to_be_building",
        "to_be_validation",
        "functional_review",
        "waiting_approval",
    ]
    for next_status in transitions:
        response = client.post(
            f"/api/v1/lpns/versions/{version_id}/transition",
            headers=headers,
            json={"to_status": next_status},
        )
        if response.status_code != 200:
            validation = client.post(
                f"/api/v1/lpns/versions/{version_id}/validate", headers=headers
            )
            pytest.fail(f"{next_status}: {response.text}; {validation.json()}")

    decision = client.post(
        f"/api/v1/lpns/approval/{approval.json()['id']}/decision",
        headers=headers,
        json={"decision": "approved"},
    )
    assert decision.status_code == 201
    approved = client.post(
        f"/api/v1/lpns/versions/{version_id}/transition",
        headers=headers,
        json={"to_status": "approved"},
    )
    assert approved.status_code == 200

    immutable = client.put(
        f"/api/v1/lpns/versions/{version_id}/content/{item_ids['storytelling']}",
        headers=headers,
        json={
            "kind": "storytelling",
            "code": "ST-001",
            "title": "Tentativa de alteração",
            "payload": {},
        },
    )
    assert immutable.status_code == 409

    documents = client.post(
        f"/api/v1/lpns/versions/{version_id}/documents",
        headers=headers,
        json={"formats": ["json", "docx", "pdf"]},
    )
    assert documents.status_code == 201, documents.text
    assert {item["content_type"] for item in documents.json()} == {
        "application/json",
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    docx_metadata = next(
        item for item in documents.json() if item["filename"].endswith(".docx")
    )
    downloaded = client.get(
        f"/api/v1/lpns/documents/{docx_metadata['id']}", headers=headers
    )
    generated_docx = Document(io.BytesIO(downloaded.content))
    document_text = "\n".join(paragraph.text for paragraph in generated_docx.paragraphs)
    cover_text = " ".join(
        node.text or "" for node in generated_docx.paragraphs[0]._p.xpath(".//w:t")
    )
    assert "Cooperativa Exemplo" in cover_text
    assert "238862" in cover_text
    assert "Automatizar recebimento" in cover_text
    assert "Vaccaro" not in cover_text
    assert "DETALHAMENTO DO PROCESSO ATUAL" in document_text
    assert "OBJETIVO E RESULTADOS ESPERADOS" in document_text
    assert "DETALHAMENTOS DO PROCESSO PROPOSTO" in document_text
    assert "APROVAÇÃO/ACEITE" in document_text
    assert generated_docx.tables[-1].cell(1, 0).text == "Admin LPN"

    cloned = client.post(
        f"/api/v1/lpns/versions/{version_id}/clone",
        headers=headers,
        json={"change_summary": "Ajustar exceção fiscal"},
    )
    assert cloned.status_code == 201
    assert cloned.json()["version_number"] == 2
    cloned_content = client.get(
        f"/api/v1/lpns/versions/{cloned.json()['id']}/content", headers=headers
    )
    assert len(cloned_content.json()) == 5


def test_ai_compose_creates_six_editable_blocks(client: TestClient) -> None:
    headers, _ = authenticate(client)
    organization = client.get("/api/v1/organizations", headers=headers).json()[0]
    headers["X-Organization-ID"] = organization["id"]
    created_client = client.post(
        "/api/v1/organizations/clients",
        headers=headers,
        json={"name": "Cliente Editor Visual"},
    )
    demand = client.post(
        "/api/v1/lpns/demands",
        headers=headers,
        json={
            "client_id": created_client.json()["id"],
            "title": "Simplificar relatório de carregamento",
            "business_area": "Expedição",
            "business_process": "Carregamento",
            "system_product": "VPE020",
            "requester_name": "Responsável da expedição",
            "priority": "medium",
            "discovery_date": "2026-08-26",
            "demand_type": "improvement",
        },
    )
    lpn = client.post(
        f"/api/v1/lpns/from-demand/{demand.json()['id']}", headers=headers
    ).json()
    version_id = lpn["current_version"]["id"]

    composed = client.post(
        f"/api/v1/lpns/versions/{version_id}/ai/compose",
        headers=headers,
        json={
            "as_is": (
                "Hoje a expedição preenche muitos filtros e confere manualmente "
                "as informações do relatório de carregamento."
            ),
            "to_be": (
                "A tela deve destacar filtros obrigatórios, validar os dados e "
                "mostrar um resumo antes de gerar o relatório."
            ),
            "constraints": "Manter regras, permissões e origem dos dados.",
        },
    )
    assert composed.status_code == 200, composed.text
    suggestion_ids = composed.json()["suggestion_ids"]
    assert len(suggestion_ids) == 6

    for suggestion_id in suggestion_ids:
        accepted = client.post(
            f"/api/v1/lpns/ai/suggestions/{suggestion_id}/decision",
            headers=headers,
            json={"decision": "accepted"},
        )
        assert accepted.status_code == 201, accepted.text

    content = client.get(
        f"/api/v1/lpns/versions/{version_id}/content", headers=headers
    ).json()
    assert len(content) == 6
    assert {item["kind"] for item in content} == {
        "storytelling",
        "objective",
        "requirement",
        "constraint",
        "pending_issue",
        "acceptance_criterion",
    }
    requirement = next(item for item in content if item["kind"] == "requirement")
    assert len(requirement["payload"]["process_steps"]) >= 2
