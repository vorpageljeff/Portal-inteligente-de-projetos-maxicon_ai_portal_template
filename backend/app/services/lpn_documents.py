import hashlib
import html
import io
import json
from collections import defaultdict
from datetime import datetime

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    Image as PdfImage,
)
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.models.lpn import (
    ApprovalDecision,
    AttachmentVersion,
    Client,
    ContentKind,
    Demand,
    DocumentGenerationJob,
    DocumentStatus,
    Evidence,
    GeneratedDocument,
    Lpn,
    LpnContentItem,
    LpnVersion,
    ProcessDiagram,
)
from app.models.security import User

SECTION_LABELS = {
    ContentKind.STORYTELLING: "DETALHAMENTO DO PROCESSO ATUAL",
    ContentKind.OBJECTIVE: "OBJETIVO E RESULTADOS ESPERADOS",
    ContentKind.REQUIREMENT: "DETALHAMENTOS DO PROCESSO PROPOSTO",
    ContentKind.CONSTRAINT: "RESTRIÇÕES/IMPEDITIVOS",
    ContentKind.PENDING_ISSUE: "INFORMAÇÕES COMPLEMENTARES",
    ContentKind.ACCEPTANCE_CRITERION: "CRITÉRIOS DE ACEITE",
}

SECTION_ORDER = tuple(SECTION_LABELS.values())
MAXICON_BLUE = RGBColor(31, 90, 166)
APPROVAL_LABELS = {
    "approved": "Aprovado",
    "rejected": "Rejeitado",
    "changes_requested": "Alterações solicitadas",
}


def _document_data(db: Session, version: LpnVersion) -> dict:
    lpn = db.get(Lpn, version.lpn_id)
    assert lpn is not None
    demand = db.get(Demand, lpn.demand_id)
    assert demand is not None
    client = db.get(Client, demand.client_id)
    analyst = db.get(User, demand.analyst_user_id)
    items = list(
        db.scalars(
            select(LpnContentItem)
            .where(LpnContentItem.lpn_version_id == version.id)
            .order_by(LpnContentItem.kind, LpnContentItem.sort_order, LpnContentItem.code)
        )
    )
    grouped: dict[str, list[dict]] = defaultdict(list)
    for item in items:
        section = SECTION_LABELS.get(item.kind)
        if not section:
            continue
        grouped[section].append(
            {"code": item.code, "title": item.title, "content": item.payload}
        )
    diagrams = list(
        db.scalars(select(ProcessDiagram).where(ProcessDiagram.lpn_version_id == version.id))
    )
    evidences = list(
        db.execute(
            select(Evidence, AttachmentVersion)
            .join(
                AttachmentVersion,
                AttachmentVersion.id == Evidence.attachment_version_id,
            )
            .where(Evidence.lpn_version_id == version.id)
            .order_by(Evidence.created_at)
        )
    )
    approvals = list(
        db.execute(
            select(ApprovalDecision, User)
            .join(User, User.id == ApprovalDecision.user_id)
            .where(ApprovalDecision.lpn_version_id == version.id)
            .order_by(ApprovalDecision.created_at)
        )
    )
    return {
        "lpn_id": str(lpn.id),
        "version": version.version_number,
        "status": version.status.value,
        "demand": {
            "client": client.name if client else "",
            "external_number": demand.external_number or "",
            "title": demand.title,
            "business_area": demand.business_area,
            "business_process": demand.business_process,
            "system_product": demand.system_product,
            "management_assignment": demand.requester_name,
            "product_owner": demand.product_owner_name or "",
            "business_analyst": analyst.full_name if analyst else "",
            "priority": demand.priority.value,
        },
        "sections": {section: grouped.get(section, []) for section in SECTION_ORDER},
        "processes": [
            {"type": diagram.process_type.value, "name": diagram.name, "model": diagram.model}
            for diagram in diagrams
        ],
        "evidences": [
            {
                "filename": attachment.filename,
                "content_type": attachment.content_type,
                "description": evidence.description or attachment.filename,
                "content": attachment.content,
            }
            for evidence, attachment in evidences
        ],
        "approvals": [
            {
                "name": user.full_name,
                "decision": decision.decision.value,
                "date": decision.created_at.isoformat(),
            }
            for decision, user in approvals
        ],
        "generated_at": datetime.utcnow().isoformat(),
    }


def _plain_text(value: object) -> str:
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        description = value.get("description")
        if description:
            return str(description)
        labels = {
            "actor": "Ator",
            "current_situation": "Situação atual",
            "problem": "Problema",
            "consequence": "Consequência",
            "desired_result": "Resultado esperado",
        }
        parts = [
            f"{labels.get(key, key.replace('_', ' ').title())}: {content}"
            for key, content in value.items()
            if content not in (None, "", [], {})
        ]
        return "\n".join(parts)
    return json.dumps(value, ensure_ascii=False, indent=2)


def _shade_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def _configure_docx(document: DocumentObject) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = MAXICON_BLUE
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def _add_document_header_footer(document: DocumentObject, data: dict) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("MAXICON SISTEMAS  |  LEVANTAMENTO DE PROCESSOS DE NEGÓCIO")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 104, 116)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Solicitação {data['demand']['external_number'] or 'não informada'}  •  "
        f"Versão {data['version']}"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 104, 116)


def _add_cover(document: DocumentObject, data: dict) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(105)
    brand = document.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand_run = brand.add_run("maxicon")
    brand_run.font.name = "Arial"
    brand_run.font.size = Pt(34)
    brand_run.font.bold = True
    brand_run.font.color.rgb = MAXICON_BLUE
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("SISTEMAS")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(50, 50, 50)
    document.add_paragraph().paragraph_format.space_after = Pt(105)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LEVANTAMENTO DE PROCESSOS DE NEGÓCIO")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = MAXICON_BLUE
    process = document.add_paragraph()
    process.alignment = WD_ALIGN_PARAGRAPH.CENTER
    process_run = process.add_run(data["demand"]["title"])
    process_run.font.size = Pt(15)
    process_run.font.bold = True
    process_run.font.color.rgb = RGBColor(55, 65, 72)
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.add_run(
        f"Cliente: {data['demand']['client']}  |  "
        f"Solicitação: {data['demand']['external_number'] or 'não informada'}"
    )
    document.add_page_break()


def _add_general_data(document: DocumentObject, data: dict) -> None:
    document.add_heading("DADOS GERAIS – LEVANTAMENTO DE PROCESSOS DE NEGÓCIO", level=1)
    demand = data["demand"]
    rows = [
        ("Cliente", demand["client"]),
        ("Solicitação", demand["external_number"] or "Não informada"),
        ("Módulos envolvidos", demand["system_product"]),
        ("Processo", demand["title"]),
        ("Área de negócio", demand["business_area"]),
        ("Product Owner", demand["product_owner"] or "Não informado"),
        ("Analista de Negócios", demand["business_analyst"] or "Não informado"),
        ("Designação de gerenciamento", demand["management_assignment"]),
    ]
    table = document.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        label_cell, value_cell = table.rows[index].cells
        _set_cell_width(label_cell, 2500)
        _set_cell_width(value_cell, 7420)
        label_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        value_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _shade_cell(label_cell, "E8EEF5")
        label_cell.paragraphs[0].add_run(label).bold = True
        value_cell.paragraphs[0].add_run(str(value))


def _add_section_content(
    document: DocumentObject, title: str, items: list[dict]
) -> None:
    document.add_heading(title, level=1)
    if not items:
        document.add_paragraph("Não informado.")
        return
    for index, item in enumerate(items):
        if len(items) > 1 or item["title"].strip().lower() not in {
            title.lower(),
            "conteúdo",
        }:
            heading = document.add_paragraph()
            heading.paragraph_format.space_after = Pt(3)
            run = heading.add_run(item["title"])
            run.bold = True
            run.font.color.rgb = RGBColor(45, 58, 68)
        text = _plain_text(item["content"])
        for paragraph in text.splitlines() or [text]:
            document.add_paragraph(paragraph)
        if index < len(items) - 1:
            document.add_paragraph()


def _add_evidences(document: DocumentObject, data: dict) -> None:
    image_evidences = [
        evidence
        for evidence in data["evidences"]
        if evidence["content_type"] in {"image/png", "image/jpeg"}
    ]
    if not image_evidences:
        return
    for evidence in image_evidences:
        caption = document.add_paragraph()
        caption.paragraph_format.space_before = Pt(6)
        caption.paragraph_format.space_after = Pt(4)
        caption.add_run(evidence["description"]).italic = True
        try:
            document.add_picture(io.BytesIO(evidence["content"]), width=Inches(6.5))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except (ValueError, TypeError):
            document.add_paragraph(f"Anexo: {evidence['filename']}")


def _add_process_diagram(document: DocumentObject, data: dict) -> None:
    document.add_heading("DIAGRAMA DO PROCESSO", level=1)
    process = next(
        (item for item in data["processes"] if item["type"] == "to_be"),
        data["processes"][0] if data["processes"] else None,
    )
    nodes = process["model"].get("nodes", []) if process else []
    if not nodes:
        document.add_paragraph("Diagrama não informado.")
        return
    table = document.add_table(rows=1, cols=len(nodes))
    table.autofit = False
    width = max(1100, 9920 // len(nodes))
    fills = ("DCEAF7", "E5F1DD", "FFF0C7", "FBE3D5")
    for index, node in enumerate(nodes):
        cell = table.cell(0, index)
        _set_cell_width(cell, width)
        _shade_cell(cell, fills[index % len(fills)])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label = str(node.get("name") or node.get("label") or node.get("type") or "Etapa")
        run = paragraph.add_run(label)
        run.bold = True
        if index < len(nodes) - 1:
            paragraph.add_run("  →")


def _build_docx(data: dict) -> bytes:
    document = Document()
    _configure_docx(document)
    _add_document_header_footer(document, data)
    _add_cover(document, data)
    _add_general_data(document, data)
    _add_section_content(
        document,
        "DETALHAMENTO DO PROCESSO ATUAL",
        data["sections"]["DETALHAMENTO DO PROCESSO ATUAL"],
    )
    _add_evidences(document, data)
    _add_section_content(
        document,
        "OBJETIVO E RESULTADOS ESPERADOS",
        data["sections"]["OBJETIVO E RESULTADOS ESPERADOS"],
    )
    _add_process_diagram(document, data)
    for section in (
        "DETALHAMENTOS DO PROCESSO PROPOSTO",
        "RESTRIÇÕES/IMPEDITIVOS",
        "INFORMAÇÕES COMPLEMENTARES",
        "CRITÉRIOS DE ACEITE",
    ):
        _add_section_content(document, section, data["sections"][section])
    document.add_heading("APROVAÇÃO/ACEITE", level=1)
    document.add_paragraph(
        "Estou de acordo com os processos descritos neste documento e ciente de que "
        "faz parte do escopo desta demanda exclusivamente o que está discriminado acima."
    )
    document.add_paragraph(
        "Qualquer alteração adicional será considerada uma nova demanda e deverá ser "
        "tratada em outra especificação ou solicitação."
    )
    approval_table = document.add_table(rows=1, cols=3)
    approval_table.style = "Table Grid"
    for index, label in enumerate(("Nome", "Decisão", "Data")):
        _shade_cell(approval_table.cell(0, index), "E8EEF5")
        approval_table.cell(0, index).paragraphs[0].add_run(label).bold = True
    for approval in data["approvals"]:
        cells = approval_table.add_row().cells
        cells[0].text = approval["name"]
        cells[1].text = APPROVAL_LABELS.get(approval["decision"], approval["decision"])
        cells[2].text = approval["date"][:10]
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _build_pdf(data: dict) -> bytes:
    output = io.BytesIO()
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="LpnHeading",
            parent=styles["Heading1"],
            textColor=colors.HexColor("#1F5AA6"),
            fontSize=14,
            leading=17,
            spaceBefore=14,
            spaceAfter=7,
        )
    )
    story = [
        Spacer(1, 55 * mm),
        Paragraph("MAXICON SISTEMAS", styles["Title"]),
        Spacer(1, 35 * mm),
        Paragraph("LEVANTAMENTO DE PROCESSOS DE NEGÓCIO", styles["Title"]),
        Paragraph(html.escape(data["demand"]["title"]), styles["Heading2"]),
        Spacer(1, 8),
        Paragraph(
            f"Cliente: {html.escape(data['demand']['client'])} | "
            f"Solicitação: {html.escape(data['demand']['external_number'] or 'não informada')}",
            styles["BodyText"],
        ),
        PageBreak(),
        Paragraph("DADOS GERAIS – LEVANTAMENTO DE PROCESSOS DE NEGÓCIO", styles["LpnHeading"]),
    ]
    demand = data["demand"]
    general_rows = [
        ["Cliente", demand["client"]],
        ["Solicitação", demand["external_number"] or "Não informada"],
        ["Módulos envolvidos", demand["system_product"]],
        ["Processo", demand["title"]],
        ["Área de negócio", demand["business_area"]],
        ["Product Owner", demand["product_owner"] or "Não informado"],
        ["Analista de Negócios", demand["business_analyst"] or "Não informado"],
        ["Designação de gerenciamento", demand["management_assignment"]],
    ]
    general_table = Table(general_rows, colWidths=[45 * mm, 125 * mm])
    general_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C5D1")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(general_table)
    for section in SECTION_ORDER:
        items = data["sections"][section]
        story.extend([Spacer(1, 12), Paragraph(section, styles["LpnHeading"])])
        for item in items:
            story.append(Paragraph(html.escape(item["title"]), styles["Heading2"]))
            safe_content = html.escape(_plain_text(item["content"]))
            story.append(Paragraph(safe_content.replace("\n", "<br/>"), styles["BodyText"]))
        if section == "DETALHAMENTO DO PROCESSO ATUAL":
            for evidence in data["evidences"]:
                if evidence["content_type"] not in {"image/png", "image/jpeg"}:
                    continue
                story.append(Paragraph(html.escape(evidence["description"]), styles["BodyText"]))
                try:
                    image = PdfImage(io.BytesIO(evidence["content"]))
                    image.drawWidth = 170 * mm
                    image.drawHeight = image.imageHeight * image.drawWidth / image.imageWidth
                    story.append(image)
                except (ValueError, TypeError):
                    story.append(Paragraph(f"Anexo: {evidence['filename']}", styles["BodyText"]))
        if section == "OBJETIVO E RESULTADOS ESPERADOS":
            process = next(
                (item for item in data["processes"] if item["type"] == "to_be"),
                data["processes"][0] if data["processes"] else None,
            )
            story.append(Paragraph("DIAGRAMA DO PROCESSO", styles["LpnHeading"]))
            nodes = process["model"].get("nodes", []) if process else []
            if nodes:
                labels = [
                    html.escape(str(node.get("name") or node.get("label") or "Etapa"))
                    for node in nodes
                ]
                story.append(Paragraph(" &nbsp; → &nbsp; ".join(labels), styles["BodyText"]))
            else:
                story.append(Paragraph("Diagrama não informado.", styles["BodyText"]))
    story.append(Paragraph("APROVAÇÃO/ACEITE", styles["LpnHeading"]))
    story.append(
        Paragraph(
            "Estou de acordo com os processos descritos neste documento e ciente de que "
            "faz parte do escopo desta demanda exclusivamente o que está discriminado acima.",
            styles["BodyText"],
        )
    )
    approval_rows = [["Nome", "Decisão", "Data"]] + [
        [
            item["name"],
            APPROVAL_LABELS.get(item["decision"], item["decision"]),
            item["date"][:10],
        ]
        for item in data["approvals"]
    ]
    approval_table = Table(approval_rows, colWidths=[85 * mm, 45 * mm, 40 * mm])
    approval_table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B8C5D1")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8EEF5")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.append(approval_table)
    SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=22 * mm,
        bottomMargin=20 * mm,
    ).build(story)
    return output.getvalue()


def _build_diagram_svg(process: dict) -> bytes:
    model = process.get("model", {})
    lanes = model.get("lanes", []) if isinstance(model, dict) else []
    nodes = model.get("nodes", []) if isinstance(model, dict) else []
    edges = model.get("edges", []) if isinstance(model, dict) else []
    lane_height = 170
    width = max(1000, 260 + len(nodes) * 190)
    height = max(220, 60 + max(len(lanes), 1) * lane_height)
    lane_positions = {
        lane.get("id"): index for index, lane in enumerate(lanes) if isinstance(lane, dict)
    }
    node_positions: dict[object, tuple[int, int]] = {}
    lane_counts: dict[object, int] = defaultdict(int)
    for node in nodes:
        lane_id = node.get("lane_id")
        lane_index = lane_positions.get(lane_id, 0)
        x = 180 + lane_counts[lane_id] * 190
        y = 75 + lane_index * lane_height
        node_positions[node.get("id")] = (x, y)
        lane_counts[lane_id] += 1
    parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}">',
        "<defs><marker id=\"arrow\" markerWidth=\"10\" markerHeight=\"7\" "
        "refX=\"9\" refY=\"3.5\" orient=\"auto\"><polygon points=\"0 0, 10 3.5, 0 7\" "
        "fill=\"#1c625d\"/></marker></defs>",
        "<rect width=\"100%\" height=\"100%\" fill=\"#ffffff\"/>",
    ]
    for index, lane in enumerate(lanes or [{"id": "default", "name": "Processo"}]):
        y = 30 + index * lane_height
        name = html.escape(str(lane.get("name") or "Raia"))
        parts.append(
            f'<rect x="20" y="{y}" width="{width - 40}" height="{lane_height - 10}" '
            'rx="10" fill="#eef7f5" stroke="#9bc9c2"/>'
        )
        parts.append(
            f'<text x="36" y="{y + 28}" font-family="Arial" font-size="15" '
            f'font-weight="bold" fill="#164f4c">{name}</text>'
        )
    for edge in edges:
        source = node_positions.get(edge.get("source"))
        target = node_positions.get(edge.get("target"))
        if not source or not target:
            continue
        parts.append(
            f'<line x1="{source[0] + 140}" y1="{source[1] + 35}" '
            f'x2="{target[0]}" y2="{target[1] + 35}" stroke="#1c625d" '
            'stroke-width="2" marker-end="url(#arrow)"/>'
        )
    for node in nodes:
        x, y = node_positions[node.get("id")]
        label = html.escape(
            str(node.get("name") or node.get("label") or node.get("type") or "Etapa")
        )
        fill = "#fff3d6" if node.get("type") == "decision" else "#ffffff"
        parts.append(
            f'<rect x="{x}" y="{y}" width="140" height="70" rx="12" '
            f'fill="{fill}" stroke="#087b72" stroke-width="2"/>'
        )
        parts.append(
            f'<text x="{x + 70}" y="{y + 40}" text-anchor="middle" '
            f'font-family="Arial" font-size="13" fill="#173b3e">{label[:28]}</text>'
        )
    parts.append("</svg>")
    return "".join(parts).encode("utf-8")


def _json_document_data(data: dict) -> dict:
    return {
        **data,
        "evidences": [
            {key: value for key, value in evidence.items() if key != "content"}
            for evidence in data["evidences"]
        ],
    }


def generate_documents(
    db: Session,
    *,
    version: LpnVersion,
    job: DocumentGenerationJob,
) -> list[GeneratedDocument]:
    data = _document_data(db, version)
    builders = {
        "json": (
            "application/json",
            lambda: json.dumps(
                _json_document_data(data), ensure_ascii=False, indent=2
            ).encode(),
        ),
        "docx": (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            lambda: _build_docx(data),
        ),
        "pdf": ("application/pdf", lambda: _build_pdf(data)),
    }
    documents: list[GeneratedDocument] = []
    job.status = DocumentStatus.PROCESSING
    db.flush()
    try:
        for format_name in job.formats:
            if format_name == "svg":
                for process in data["processes"]:
                    content = _build_diagram_svg(process)
                    document = GeneratedDocument(
                        job_id=job.id,
                        lpn_version_id=version.id,
                        filename=(
                            f"lpn-{version.lpn_id}-v{version.version_number}-"
                            f"{process['type']}.svg"
                        ),
                        content_type="image/svg+xml",
                        content=content,
                        sha256=hashlib.sha256(content).hexdigest(),
                    )
                    db.add(document)
                    documents.append(document)
                continue
            content_type, builder = builders[format_name]
            content = builder()
            document = GeneratedDocument(
                job_id=job.id,
                lpn_version_id=version.id,
                filename=f"lpn-{version.lpn_id}-v{version.version_number}.{format_name}",
                content_type=content_type,
                content=content,
                sha256=hashlib.sha256(content).hexdigest(),
            )
            db.add(document)
            documents.append(document)
        job.status = DocumentStatus.GENERATED
        job.completed_at = datetime.utcnow()
        version.document_status = DocumentStatus.GENERATED
    except Exception as exc:
        job.status = DocumentStatus.FAILED
        job.error_message = str(exc)[:2000]
        version.document_status = DocumentStatus.FAILED
        raise
    db.flush()
    return documents
